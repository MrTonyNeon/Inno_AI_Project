import os
import re
import time
import fitz
import argparse
import docx
from tqdm import tqdm
from pathlib import Path
from openai import OpenAI
from llama_cpp import Llama
from dotenv import load_dotenv

# Load environment variables (for API keys)
load_dotenv()

# Set up command line arguments
parser = argparse.ArgumentParser(description='Anonymize student papers using AI API')
parser.add_argument('--mode', choices=['regex', 'local_ai', 'online_ai'], default='regex',
                    help='Mode for anonymization: regex (pattern matching), local_ai (local model), online_ai (OpenAI API)')
parser.add_argument('--max_files', type=int, default=2, help='Maximum number of files to process (default: 2)')
parser.add_argument('--max_tokens', type=int, default=20000, help='Maximum tokens to send to API per file')
parser.add_argument("--model", default='gpt-4.1-nano',
                    help="Models to use for evaluation (model_name)")
args = parser.parse_args()


def create_directory(directory_path):
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)
        print(f"Created directory: {directory_path}")


def extract_text_from_pdf(file_path):
    text = ""
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        print(f"PDF text extraction failed for {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)

        for para in doc.paragraphs:
            text += para.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return ""


def extract_text_from_file(file_path):
    file_extension = file_path.suffix.lower()

    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        print(f"Unsupported file type: {file_extension}")
        return ""


def create_client():
    if args.mode != 'online_ai':
        return None
    elif os.environ.get("AITUNNEL_API_KEY"):
        return OpenAI(
            api_key=os.environ["AITUNNEL_API_KEY"],
            base_url="https://api.aitunnel.ru/v1"
        )
    elif os.environ.get("OPENAI_API_KEY"):
        return OpenAI()
    else:
        raise ValueError(
            "OpenAI API key must be provided via OPENAI_API_KEY environment variable when using online_ai mode")


def anonymize_text_with_online_ai(client, text, model_name):
    if len(text) > args.max_tokens * 4:
        print(f"Text too long ({len(text)} chars), trimming to ~{args.max_tokens} tokens")
        text = text[:args.max_tokens * 4]

    system_prompt = """
    You are a text anonymization tool. Your task is to find and replace all personal names in the document.

    Replace all people's names with PERSON_N format, where N is a sequential number 
    (first name found is PERSON_1, second is PERSON_2, etc.).

    Rules:
    1. Maintain consistent numbering (the same name should always be replaced with the same PERSON_N).
    2. Don't alter any text except for replacing names.
    3. Ignore company names, organization names, and product names - only replace people's names.
    4. Return only the transformed text without any additional comments.
    """

    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            temperature=0,
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error with AI model '{model_name}': {e}")
        return text


def anonymize_text_with_local_ai(text):
    model_path = "models/mistral-7b-instruct-v0.2.Q4_K_M.gguf"

    model = Llama(
        model_path=model_path,
        n_ctx=4096,
        n_gpu_layers=-1,
        n_threads=4,
        verbose=False
    )

    prompt = f"""
    You are a text anonymization tool. Your task is to find and replace all personal names in the document.

    Replace all people's names with PERSON_N format, where N is a sequential number 
    (first name found is PERSON_1, second is PERSON_2, etc.).

    Rules:
    1. Maintain consistent numbering (the same name should always be replaced with the same PERSON_N).
    2. Don't alter any text except for replacing names.
    3. Ignore company names, organization names, and product names - only replace people's names.
    4. Return only the transformed text without any additional comments.

    Original text:
    {text}

    Anonymized text:
    """

    response = model.create_completion(
        prompt,
        max_tokens=2560,
        temperature=0,
        stop=["</s>", "\n\n"],
        echo=False
    )

    anonymized_text = response["choices"][0]["text"].strip()
    return anonymized_text


def anonymize_text_locally(text):
    lines = text.splitlines()
    name_pattern = r'\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)+)\b'
    name_mapping = {}
    counter = 1

    check_words = ["Analysis", "Script", "Magazine", "Bank", "Team",
                   "Interview", "Report", "Review", "Summary", "Results", "Hough"]

    def replace_name(match):
        nonlocal counter
        name = match.group(1)
        if any(check_word in name for check_word in check_words):
            return name
        if name not in name_mapping:
            name_mapping[name] = f"PERSON_{counter}"
            counter += 1
        return name_mapping[name]

    anonymized_lines = []
    for line in lines:
        if re.search(r'^\d+\.\d+', line):
            anonymized_lines.append(line)
        else:
            anonymized_line = re.sub(name_pattern, replace_name, line)
            anonymized_lines.append(anonymized_line)

    anonymized_text = "\n".join(anonymized_lines)
    return anonymized_text


def anonymize_student_papers():
    base_dir = Path(".")
    student_papers_dir = base_dir / "student_papers"
    anonymized_dir = base_dir / "anonymized_papers"

    create_directory(anonymized_dir)
    client = create_client()

    team_folders = [f for f in student_papers_dir.iterdir() if f.is_dir()]
    processed_files = 0

    for team_folder in tqdm(team_folders, desc="Processing team folders"):
        team_name = team_folder.name

        document_files = list(team_folder.glob("*.pdf")) + list(team_folder.glob("*.docx"))

        if not document_files:
            print(f"No files found in {team_folder}")
            continue

        for doc_file in document_files:
            if processed_files >= args.max_files:
                print(f"Reached maximum files limit ({args.max_files}). Stopping.")
                break

            output_filename = f"{team_name}_{doc_file.stem}_anonymized.txt"
            output_path = anonymized_dir / output_filename

            if output_path.exists():
                print(f"Skipping already processed file: {output_path}")
                continue

            print(f"Processing: {doc_file}")
            text = extract_text_from_file(doc_file)

            if not text.strip():
                print(f"Warning: No text extracted from {doc_file}")
                continue

            start_time = time.time()

            if args.mode == 'online_ai':
                anonymized_text = anonymize_text_with_online_ai(client, text, args.model)
            elif args.mode == 'local_ai':
                anonymized_text = anonymize_text_with_local_ai(text)
            else:
                anonymized_text = anonymize_text_locally(text)

            end_time = time.time()

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(anonymized_text)

            processing_time = end_time - start_time
            print(f"Saved anonymized text to {output_path} (processing time: {processing_time:.2f} seconds)")

            processed_files += 1

            if client:
                time.sleep(1)

        if processed_files >= args.max_files:
            break

    return processed_files


if __name__ == "__main__":
    print("Starting anonymization process with the following settings:")
    print(f"- Anonymization mode: {args.mode}")
    print(f"- Maximum files to process: {args.max_files}")

    processed = anonymize_student_papers()
    print(f"Anonymization complete! Processed {processed} files.")